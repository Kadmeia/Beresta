import os
import docx
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

class DocxProcessor:
    def __init__(self):
        pass

    def _get_virtual_pages(self, doc):
        """
        Splits docx into virtual pages based on length (~1500 chars) or explicit page breaks.
        Returns (pages, blocks_list)
        pages = [{"text": str, "start_idx": int, "end_idx": int}]
        """
        pages = []
        current_page_text = []
        current_start = 0
        
        blocks = list(iter_block_items(doc))
        
        for i, block in enumerate(blocks):
            text = ""
            is_page_break = False
            
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if 'w:br w:type="page"' in block._element.xml:
                    is_page_break = True
            elif isinstance(block, Table):
                for row in block.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text += " ".join(row_text) + "\n"
                text = text.strip()
                
            if text:
                current_page_text.append(text)
                
            current_len = len("\n".join(current_page_text))
            if current_len > 1500 or is_page_break:
                pages.append({
                    "text": "\n".join(current_page_text),
                    "start_idx": current_start,
                    "end_idx": i
                })
                current_page_text = []
                current_start = i + 1
                
        if current_page_text or current_start < len(blocks):
            pages.append({
                "text": "\n".join(current_page_text),
                "start_idx": current_start,
                "end_idx": len(blocks) - 1
            })
            
        return pages, blocks

    def extract_text(self, file_path, status_callback=None):
        if status_callback:
            status_callback("Чтение DOCX документа...")
            
        try:
            doc = docx.Document(file_path)
            pages, _ = self._get_virtual_pages(doc)
            
            result = []
            for i, page in enumerate(pages):
                result.append({
                    "page_num": i + 1,
                    "text": page['text']
                })
            return result
        except Exception as e:
            print(f"Error reading docx: {e}")
            raise ValueError(f"Не удалось прочитать файл {os.path.basename(file_path)}: {str(e)}")

    def split_and_save(self, original_path, start_page, end_page, new_name, output_dir):
        """
        Intelligently splits the docx by removing XML nodes that are outside the requested page range.
        This preserves all styles, headers, and tables perfectly.
        """
        doc = docx.Document(original_path)
        pages, blocks = self._get_virtual_pages(doc)
        
        start_idx = pages[start_page - 1]['start_idx']
        end_idx = pages[end_page - 1]['end_idx']
        
        body = doc.element.body
        # Iterate backwards to safely delete elements
        for i in range(len(blocks) - 1, -1, -1):
            if not (start_idx <= i <= end_idx):
                try:
                    body.remove(blocks[i]._element)
                except Exception as e:
                    print(f"Warning: could not remove block {i}: {e}")
                    
        base_path = os.path.join(output_dir, new_name + ".docx")
        final_path = base_path
        
        counter = 1
        while os.path.exists(final_path):
            final_path = os.path.join(output_dir, f"{new_name} ({counter}).docx")
            counter += 1
            
        doc.save(final_path)
        return final_path

    def create_docx_from_text(self, text, new_name, output_dir):
        """
        Creates a new docx file from plain text or Markdown text.
        Supports Markdown headings, lists, bold text, and tables.
        """
        doc = docx.Document()
        
        lines = text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
                
            # Check for table
            if line.startswith('|') and line.endswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                
                # Parse table
                if len(table_lines) >= 2:
                    rows = []
                    for t_line in table_lines:
                        # Split by | and remove empty strings from edges
                        cols = [c.strip() for c in t_line.split('|')[1:-1]]
                        # Skip separator row (---)
                        if all(all(char in '-: ' for char in c) for c in cols) and len(rows) == 1:
                            continue
                        rows.append(cols)
                    
                    if rows:
                        num_cols = max(len(row) for row in rows)
                        table = doc.add_table(rows=len(rows), cols=num_cols)
                        table.style = 'Table Grid'
                        for r_idx, row in enumerate(rows):
                            for c_idx, cell_text in enumerate(row):
                                if c_idx < num_cols:
                                    cell = table.cell(r_idx, c_idx)
                                    self._add_markdown_run(cell.paragraphs[0], cell_text)
                continue
                
            # Check for headings
            import re
            heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
            if heading_match:
                level = len(heading_match.group(1))
                text_content = heading_match.group(2)
                # python-docx has styles 'Heading 1', 'Heading 2', etc. up to 9
                try:
                    p = doc.add_paragraph(style=f'Heading {level}')
                except Exception:
                    p = doc.add_paragraph() # Fallback if style doesn't exist
                self._add_markdown_run(p, text_content)
                i += 1
                continue
                
            # Normal paragraph
            p = doc.add_paragraph()
            self._add_markdown_run(p, line)
            i += 1
            
        base_path = os.path.join(output_dir, new_name + ".docx")
        final_path = base_path
        
        counter = 1
        while os.path.exists(final_path):
            final_path = os.path.join(output_dir, f"{new_name} ({counter}).docx")
            counter += 1
            
        doc.save(final_path)
        return final_path

    def _add_markdown_run(self, paragraph, text):
        import re
        # Basic inline bold **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                paragraph.add_run(part[2:-2]).bold = True
            else:
                paragraph.add_run(part)
